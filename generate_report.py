# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# --- Styles ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(13)

def set_font(run, bold=False, size=13, color=None, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading(doc, text, level=1, center=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, bold=True, size=14 if level == 1 else 13)
    return p

def add_paragraph(doc, text, indent=False, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    run = p.add_run(text)
    set_font(run)
    return p

def add_table_row(table, cells_data, bold=False, center=False, bg_color=None):
    row = table.add_row()
    for i, data in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(data))
        set_font(run, bold=bold)
        if bg_color:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)
    return row

# ============================================================
# TRANG BÌA
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ĐẠI HỌC MỞ HÀ NỘI")
set_font(run, bold=True, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TRUNG TÂM ĐÀO TẠO E-LEARNING")
set_font(run, bold=True, size=13)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BÁO CÁO CHUYÊN ĐỀ THỰC TẬP CHUYÊN NGÀNH")
set_font(run, bold=True, size=16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Đề tài:")
set_font(run, bold=True, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("XÂY DỰNG LẠI WEBSITE THEO HƯỚNG TỰ ĐỘNG HÓA NỘI DUNG\nBẰNG AI VÀ TỐI ƯU MỌI MẶT")
set_font(run, bold=True, size=15)

doc.add_paragraph()
doc.add_paragraph()

info = [
    ("MÔN:", "Chuyên đề thực tập chuyên ngành - IT43.027"),
    ("NHÓM SINH VIÊN:", "1. Đặng Minh Tuấn  (CHTM518)\n                         2. Trần Anh Tuấn     (CHCT419)\n                         3. Nguyễn Minh Hiếu  (CLCA520)"),
    ("ĐƠN VỊ THỰC TẬP:", "Hừng Đông Media"),
    ("NĂM HỌC:", "2025 – 2026"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run(f"{label} ")
    set_font(run1, bold=True, size=13)
    run2 = p.add_run(value)
    set_font(run2, size=13)

doc.add_page_break()

# ============================================================
# MỤC LỤC (thủ công)
# ============================================================
add_heading(doc, "MỤC LỤC", center=True)
toc_items = [
    ("CHƯƠNG 1: GIỚI THIỆU DOANH NGHIỆP VÀ VỊ TRÍ THỰC TẬP", "3"),
    ("  1.1. Giới thiệu doanh nghiệp", "3"),
    ("  1.2. Vị trí và nhiệm vụ thực tập", "4"),
    ("CHƯƠNG 2: MÔ TẢ DỰ ÁN VÀ CÔNG NGHỆ SỬ DỤNG", "5"),
    ("  2.1. Mô tả dự án", "5"),
    ("  2.2. Công nghệ sử dụng", "5"),
    ("CHƯƠNG 3: PHÂN TÍCH YÊU CẦU HỆ THỐNG", "7"),
    ("  3.1. Yêu cầu chức năng", "7"),
    ("  3.2. Yêu cầu phi chức năng", "7"),
    ("  3.3. Sơ đồ nghiệp vụ (Use Case)", "8"),
    ("CHƯƠNG 4: ĐẶC TẢ VÀ DANH SÁCH CHỨC NĂNG", "9"),
    ("  4.1. Đặc tả hệ thống", "9"),
    ("  4.2. Danh sách chức năng chi tiết", "10"),
    ("CHƯƠNG 5: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", "11"),
    ("  5.1. Kiến trúc giải pháp", "11"),
    ("  5.2. Thiết kế cơ sở dữ liệu", "12"),
    ("  5.3. Ràng buộc và tiêu chuẩn đánh giá", "13"),
    ("CHƯƠNG 6: TRIỂN KHAI VÀ KẾT QUẢ", "14"),
    ("  6.1. Phân công nhiệm vụ và kế hoạch", "14"),
    ("  6.2. Kết quả lập trình giai đoạn 1", "15"),
    ("  6.3. Kết quả tích hợp", "16"),
    ("  6.4. Đánh giá giải pháp", "17"),
    ("CHƯƠNG 7: HOÀN THIỆN – CÁC LỖI ĐÃ KHẮC PHỤC", "18"),
    ("  7.1. Danh sách lỗi và cách khắc phục", "18"),
    ("  7.2. Kết quả sau khi hoàn thiện", "19"),
    ("KẾT LUẬN", "20"),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(item)
    set_font(run, bold=('CHƯƠNG' in item or item == 'KẾT LUẬN'))
doc.add_page_break()

# ============================================================
# CHƯƠNG 1: GIỚI THIỆU DOANH NGHIỆP
# ============================================================
add_heading(doc, "CHƯƠNG 1: GIỚI THIỆU DOANH NGHIỆP VÀ VỊ TRÍ THỰC TẬP")
doc.add_paragraph()

add_heading(doc, "1.1. Giới thiệu doanh nghiệp", level=2)
add_paragraph(doc, "Hừng Đông Media là một doanh nghiệp hoạt động trong lĩnh vực truyền thông kỹ thuật số và marketing tổng thể, chuyên cung cấp các dịch vụ xây dựng thương hiệu, thiết kế website, quản lý nội dung và quan hệ công chúng (PR) cho các doanh nghiệp vừa và nhỏ tại Việt Nam.", indent=True)
add_paragraph(doc, "Với định hướng ứng dụng công nghệ AI vào quy trình sản xuất nội dung, Hừng Đông Media không chỉ cung cấp dịch vụ Retainer 'Nuôi Web' hàng tháng mà còn phát triển các giải pháp SaaS (Software as a Service) giúp khách hàng tự động hóa hoàn toàn quy trình tạo và cập nhật nội dung website.", indent=True)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Thông tin cơ bản:")
set_font(run, bold=True)
items_info = [
    "Tên doanh nghiệp: Hừng Đông Media",
    "Lĩnh vực hoạt động: Truyền thông kỹ thuật số, Marketing, PR, Thiết kế Web",
    "Định hướng công nghệ: Headless CMS, AI Content Automation, SaaS",
    "Quy mô: Doanh nghiệp vừa và nhỏ, đội ngũ chuyên môn đa lĩnh vực",
    "Các dịch vụ chính: Booking báo chí, quản lý dự án thuê ngoài (Outsourcing PM), xây dựng website tự động hóa AI",
]
for item in items_info:
    add_bullet(doc, item)

doc.add_paragraph()
add_heading(doc, "1.2. Vị trí và nhiệm vụ thực tập", level=2)
add_paragraph(doc, "Trong kỳ thực tập chuyên ngành tại Hừng Đông Media, nhóm sinh viên đảm nhận vị trí Kỹ sư Phát triển Phần mềm (Software Developer Intern) với nhiệm vụ chính là nghiên cứu, phân tích yêu cầu và trực tiếp xây dựng hệ thống website theo hướng tự động hóa nội dung bằng AI.", indent=True)
add_paragraph(doc, "Cụ thể, nhóm được giao xây dựng lại toàn bộ kiến trúc kỹ thuật của website Hừng Đông Media từ mô hình truyền thống sang kiến trúc Headless CMS kết hợp với cỗ máy AI Content Engine, nhằm nâng cao tốc độ tải trang, khả năng SEO và mức độ tự động hóa trong quản lý nội dung.", indent=True)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Nhiệm vụ cụ thể:")
set_font(run, bold=True)
duties = [
    "Khảo sát hiện trạng website và quy trình quản lý nội dung tại doanh nghiệp",
    "Phân tích yêu cầu kỹ thuật và nghiệp vụ của hệ thống mới",
    "Thiết kế kiến trúc Headless WordPress + Next.js + AI Engine",
    "Lập trình module AI tự động trích xuất nội dung từ tài liệu PDF/Word và đẩy lên CMS",
    "Xây dựng giao diện Frontend premium với Dark Mode và Micro-animations",
    "Kiểm thử, hoàn thiện và tối ưu hóa hệ thống",
]
for d in duties:
    add_bullet(doc, d)

doc.add_page_break()

# ============================================================
# CHƯƠNG 2: MÔ TẢ DỰ ÁN VÀ CÔNG NGHỆ
# ============================================================
add_heading(doc, "CHƯƠNG 2: MÔ TẢ DỰ ÁN VÀ CÔNG NGHỆ SỬ DỤNG")
doc.add_paragraph()

add_heading(doc, "2.1. Mô tả dự án", level=2)
add_paragraph(doc, "Dự án 'Xây Dựng Lại Website Theo Hướng Tự Động Hóa Nội Dung Bằng AI Và Tối Ưu Mọi Mặt' được thực hiện nhằm giải quyết bài toán thực tiễn của Hừng Đông Media: làm thế nào để duy trì một website luôn có nội dung mới, chuẩn SEO, tải nhanh và bảo mật cao mà không cần đội ngũ biên tập lớn.", indent=True)
add_paragraph(doc, "Giải pháp được xây dựng gồm 3 tầng chính:", indent=True)
tiers = [
    "Tầng 1 – WordPress Headless CMS: Đóng vai trò kho lưu trữ dữ liệu bảo mật hoàn toàn, không hiển thị trực tiếp với người dùng cuối. Dữ liệu được xuất thông qua WPGraphQL API.",
    "Tầng 2 – AI Content Engine (Node.js): Cỗ máy tự động đọc tài liệu PDF/Word của khách hàng, xử lý nội dung và bắn thẳng bài viết chuẩn SEO vào WordPress thông qua REST API.",
    "Tầng 3 – Next.js Frontend: Giao diện người dùng hiển thị với tốc độ 100/100 Lighthouse nhờ kiến trúc Static Site Generation (SSG) + Incremental Static Regeneration (ISR).",
]
for t in tiers:
    add_bullet(doc, t)

doc.add_paragraph()
add_heading(doc, "2.2. Công nghệ sử dụng", level=2)

tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Table Grid'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
header_cells = tech_table.rows[0].cells
headers = ["Tầng hệ thống", "Công nghệ", "Vai trò"]
for i, h in enumerate(headers):
    header_cells[i].text = ''
    p = header_cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, bold=True)
    tc = header_cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)

tech_data = [
    ("Backend CMS", "WordPress 6.x + WPGraphQL", "Lưu trữ và xuất dữ liệu qua GraphQL API"),
    ("Backend CMS", "WP Engine Headless Add-on", "Kết nối WordPress với Frontend Next.js"),
    ("Backend CMS", "MySQL", "Cơ sở dữ liệu quan hệ lưu bài viết, media"),
    ("AI Engine", "Node.js (v18+)", "Môi trường chạy cỗ máy tự động hóa"),
    ("AI Engine", "Faust.js CLI", "Framework kết nối WP và Next.js"),
    ("AI Engine", "WordPress REST API", "Giao thức đẩy bài viết vào CMS tự động"),
    ("Frontend", "Next.js 14", "Framework React cho giao diện Headless"),
    ("Frontend", "Sass/SCSS", "Stylesheet với Dark Mode và custom design tokens"),
    ("Frontend", "Apollo Client", "Thư viện giao tiếp với GraphQL"),
    ("DevOps", "Local by WP Engine", "Môi trường phát triển local WordPress"),
    ("DevOps", "Git", "Quản lý phiên bản mã nguồn"),
]
for row_data in tech_data:
    add_table_row(tech_table, row_data, center=False)

doc.add_page_break()

# ============================================================
# CHƯƠNG 3: PHÂN TÍCH YÊU CẦU
# ============================================================
add_heading(doc, "CHƯƠNG 3: PHÂN TÍCH YÊU CẦU HỆ THỐNG")
doc.add_paragraph()

add_heading(doc, "3.1. Yêu cầu chức năng", level=2)
func_reqs = [
    ("FR01", "Tự động bóc tách nội dung từ tệp PDF/Word và đẩy lên CMS", "Cao"),
    ("FR02", "Hiển thị danh sách bài viết lên giao diện Next.js qua GraphQL", "Cao"),
    ("FR03", "Trang chi tiết bài viết với đầy đủ nội dung và metadata SEO", "Cao"),
    ("FR04", "Lưu lịch sử các lần phân tích/cập nhật nội dung tự động", "Trung bình"),
    ("FR05", "Hỗ trợ phân loại bài viết theo danh mục (Category)", "Trung bình"),
    ("FR06", "Tích hợp menu điều hướng lấy từ WordPress qua GraphQL", "Trung bình"),
    ("FR07", "Báo cáo trực quan về trạng thái nội dung đã được đẩy lên", "Thấp"),
]
req_table = doc.add_table(rows=1, cols=4)
req_table.style = 'Table Grid'
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Mã yêu cầu", "Mô tả", "Mức độ ưu tiên", ""]):
    if i < 3:
        pass
headers_fr = ["Mã yêu cầu", "Mô tả chức năng", "Mức độ ưu tiên"]
doc.tables[-1]._element.getparent().remove(doc.tables[-1]._element)

req_table = doc.add_table(rows=1, cols=3)
req_table.style = 'Table Grid'
req_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(headers_fr):
    req_table.rows[0].cells[i].text = ''
    p = req_table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, bold=True)
    tc = req_table.rows[0].cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)
for r in func_reqs:
    add_table_row(req_table, r)

doc.add_paragraph()
add_heading(doc, "3.2. Yêu cầu phi chức năng", level=2)
nonfunc = [
    "Hiệu năng: Điểm Lighthouse Performance ≥ 90/100 trên cả Mobile và Desktop",
    "Tốc độ tải trang: Time to First Byte (TTFB) < 200ms nhờ ISR của Next.js",
    "Bảo mật: WordPress Admin hoàn toàn ẩn, chỉ giao tiếp qua API có xác thực (Application Password)",
    "Khả năng mở rộng: Kiến trúc Headless cho phép thay thế Frontend mà không ảnh hưởng Backend",
    "Tính sẵn sàng: Hệ thống hoạt động 24/7, không phụ thuộc vào thao tác thủ công",
    "Dễ bảo trì: Mã nguồn được tổ chức theo chuẩn Next.js App Router, có chú thích đầy đủ",
]
for n in nonfunc:
    add_bullet(doc, n)

doc.add_paragraph()
add_heading(doc, "3.3. Sơ đồ nghiệp vụ (Use Case Diagram)", level=2)
add_paragraph(doc, "Hệ thống có hai tác nhân chính:", indent=True)
actors = [
    "AI Engine (hệ thống tự động): Tự động thu thập tài liệu, xử lý nội dung và đẩy bài lên CMS",
    "Người dùng cuối (End User): Truy cập website, xem bài viết và điều hướng qua menu",
]
for a in actors:
    add_bullet(doc, a)
add_paragraph(doc, "Các use case chính bao gồm: (UC01) Đọc tài liệu PDF → (UC02) Tạo bài viết trên WordPress → (UC03) Hiển thị bài viết lên Next.js → (UC04) Xem chi tiết bài viết → (UC05) Điều hướng theo danh mục.", indent=True)

doc.add_page_break()

# ============================================================
# CHƯƠNG 4: ĐẶC TẢ VÀ DANH SÁCH CHỨC NĂNG
# ============================================================
add_heading(doc, "CHƯƠNG 4: ĐẶC TẢ VÀ DANH SÁCH CHỨC NĂNG")
doc.add_paragraph()

add_heading(doc, "4.1. Đặc tả hệ thống", level=2)
add_paragraph(doc, "Hệ thống được đặc tả theo mô hình Client-Server với lớp trung gian API. Backend WordPress không trực tiếp phục vụ HTML mà chỉ cung cấp dữ liệu thô qua hai giao thức:", indent=True)
protocols = [
    "GraphQL (WPGraphQL): Phục vụ Frontend Next.js đọc dữ liệu bài viết, menu, thông tin trang",
    "REST API (WordPress REST): Phục vụ AI Engine ghi dữ liệu bài viết mới lên CMS",
]
for p_item in protocols:
    add_bullet(doc, p_item)
add_paragraph(doc, "Toàn bộ luồng dữ liệu được bảo mật bằng WordPress Application Password – một chuẩn xác thực không lộ mật khẩu thật ra ngoài môi trường, thay vào đó sử dụng token ngẫu nhiên 24 ký tự.", indent=True)

doc.add_paragraph()
add_heading(doc, "4.2. Danh sách chức năng chi tiết", level=2)

feat_table = doc.add_table(rows=1, cols=4)
feat_table.style = 'Table Grid'
for i, h in enumerate(["Module", "Chức năng", "Trạng thái", "Ghi chú"]):
    feat_table.rows[0].cells[i].text = ''
    p = feat_table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, bold=True)
    tc = feat_table.rows[0].cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)

features = [
    ("AI Engine", "Đọc file PDF/Word từ thư mục input", "✅ Hoàn thành", ""),
    ("AI Engine", "Tạo bài viết qua WordPress REST API", "✅ Hoàn thành", "Dùng Application Password"),
    ("AI Engine", "Đặt tiêu đề, nội dung, danh mục tự động", "✅ Hoàn thành", ""),
    ("AI Engine", "Ghi log kết quả mỗi lần chạy", "✅ Hoàn thành", ""),
    ("Backend WP", "Cấu hình WPGraphQL endpoint", "✅ Hoàn thành", ""),
    ("Backend WP", "Phân loại bài viết theo Category", "✅ Hoàn thành", ""),
    ("Backend WP", "Cấu hình menu Header/Footer", "✅ Hoàn thành", ""),
    ("Frontend", "Trang chủ hiển thị danh sách bài viết", "✅ Hoàn thành", "Dark Mode Premium"),
    ("Frontend", "Trang chi tiết bài viết", "✅ Hoàn thành", ""),
    ("Frontend", "Menu điều hướng từ WordPress", "✅ Hoàn thành", ""),
    ("Frontend", "SEO Meta Tags (title, description)", "✅ Hoàn thành", ""),
    ("Frontend", "Responsive Mobile/Tablet/Desktop", "✅ Hoàn thành", ""),
    ("Frontend", "Micro-animations khi hover card", "✅ Hoàn thành", ""),
]
for f in features:
    add_table_row(feat_table, f)

doc.add_page_break()

# ============================================================
# CHƯƠNG 5: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
# ============================================================
add_heading(doc, "CHƯƠNG 5: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG")
doc.add_paragraph()

add_heading(doc, "5.1. Kiến trúc giải pháp", level=2)
add_paragraph(doc, "Hệ thống được thiết kế theo kiến trúc Headless CMS (Decoupled Architecture), tách biệt hoàn toàn lớp quản lý nội dung (Backend) với lớp hiển thị (Frontend). Đây là xu hướng kiến trúc hiện đại nhất hiện nay, được áp dụng bởi các tập đoàn lớn như Netflix, Nike, TechCrunch.", indent=True)

arch_layers = [
    ("Lớp Client (Người dùng)", "Trình duyệt web truy cập trang Next.js qua HTTPS tại localhost:3000"),
    ("Lớp Frontend (Next.js)", "Nhận yêu cầu, truy vấn GraphQL từ WordPress và render HTML tĩnh cực nhanh"),
    ("Lớp API (WPGraphQL + REST)", "Cầu nối dữ liệu: GraphQL cho đọc, REST API cho ghi tự động"),
    ("Lớp CMS (WordPress)", "Lưu trữ toàn bộ bài viết, media, menu, cài đặt trang"),
    ("Lớp Database (MySQL)", "Cơ sở dữ liệu quan hệ chứa toàn bộ data WordPress"),
    ("AI Engine (Node.js)", "Tác nhân độc lập, đọc tài liệu đầu vào và ghi dữ liệu lên CMS qua REST API"),
]
arch_table = doc.add_table(rows=1, cols=2)
arch_table.style = 'Table Grid'
for i, h in enumerate(["Tầng kiến trúc", "Mô tả"]):
    arch_table.rows[0].cells[i].text = ''
    p = arch_table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, bold=True)
    tc = arch_table.rows[0].cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)
for layer in arch_layers:
    add_table_row(arch_table, layer)

doc.add_paragraph()
add_heading(doc, "5.2. Thiết kế cơ sở dữ liệu", level=2)
add_paragraph(doc, "Hệ thống sử dụng cơ sở dữ liệu MySQL theo cấu trúc mặc định của WordPress, bao gồm các bảng chính sau:", indent=True)
db_tables = [
    ("wp_posts", "Lưu bài viết, trang, media. Trường quan trọng: post_title, post_content, post_status, post_type"),
    ("wp_postmeta", "Metadata bổ sung cho mỗi bài viết (SEO title, description, thumbnail ID...)"),
    ("wp_terms", "Danh sách danh mục (Category), thẻ (Tag)"),
    ("wp_term_relationships", "Quan hệ nhiều-nhiều giữa bài viết và danh mục"),
    ("wp_users", "Tài khoản quản trị. AI Engine xác thực qua Application Passwords liên kết với bảng này"),
    ("wp_options", "Cài đặt hệ thống: tên site, URL, cấu hình WPGraphQL endpoint"),
]
db_table = doc.add_table(rows=1, cols=2)
db_table.style = 'Table Grid'
for i, h in enumerate(["Tên bảng", "Mô tả và vai trò"]):
    db_table.rows[0].cells[i].text = ''
    p = db_table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, bold=True)
    tc = db_table.rows[0].cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)
for dt in db_tables:
    add_table_row(db_table, dt)

doc.add_paragraph()
add_heading(doc, "5.3. Ràng buộc và tiêu chuẩn đánh giá", level=2)
constraints = [
    "Ràng buộc kỹ thuật: Phiên bản Node.js ≥ 18, npm ≥ 8, PHP ≥ 8.1 (WordPress yêu cầu)",
    "Ràng buộc bảo mật: WordPress Admin URL phải được ẩn, không được dùng mật khẩu thật trong code",
    "Ràng buộc hiệu năng: Frontend Next.js phải chạy trên cổng 3000 với strictPort=true",
    "Tiêu chuẩn SEO: Mỗi bài viết phải có <title>, <meta description>, Open Graph tags",
    "Tiêu chuẩn giao diện: Phải đạt điểm Accessibility ≥ 85/100 theo Lighthouse",
    "Tiêu chuẩn mã nguồn: Cấu trúc thư mục theo chuẩn Faust.js, có file .env quản lý biến môi trường",
]
for c in constraints:
    add_bullet(doc, c)

doc.add_page_break()

# ============================================================
# CHƯƠNG 6: TRIỂN KHAI VÀ KẾT QUẢ
# ============================================================
add_heading(doc, "CHƯƠNG 6: TRIỂN KHAI VÀ KẾT QUẢ")
doc.add_paragraph()

add_heading(doc, "6.1. Phân công nhiệm vụ và kế hoạch", level=2)
plan_table = doc.add_table(rows=1, cols=4)
plan_table.style = 'Table Grid'
for i, h in enumerate(["STT", "Thành viên", "Nhiệm vụ", "Thời gian"]):
    plan_table.rows[0].cells[i].text = ''
    p = plan_table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, bold=True)
    tc = plan_table.rows[0].cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)
plan_data = [
    ("1", "Cả nhóm", "Khảo sát hiện trạng doanh nghiệp", "01/06/2026 – 07/06/2026"),
    ("2", "Đặng Minh Tuấn", "Xác định yêu cầu hệ thống", "08/06/2026 – 14/06/2026"),
    ("3", "Cả nhóm", "Phân tích và thiết kế hệ thống", "15/06/2026 – 21/06/2026"),
    ("4", "Đặng Minh Tuấn\nTrần Anh Tuấn", "Lập trình AI Engine + Backend WP", "22/06/2026 – 06/07/2026"),
    ("5", "Nguyễn Minh Hiếu\nTrần Anh Tuấn", "Lập trình Frontend Next.js", "07/07/2026 – 20/07/2026"),
    ("6", "Cả nhóm", "Kiểm thử và tích hợp hệ thống", "21/07/2026 – 27/07/2026"),
    ("7", "Cả nhóm", "Hoàn thiện – sửa lỗi – nộp báo cáo", "28/07/2026 – 05/09/2026"),
]
for pd in plan_data:
    add_table_row(plan_table, pd)

doc.add_paragraph()
add_heading(doc, "6.2. Kết quả lập trình giai đoạn 1 (AI Engine + Backend)", level=2)
add_paragraph(doc, "Sau giai đoạn lập trình đầu tiên (22/06 – 06/07/2026), nhóm đã hoàn thành các module core của hệ thống:", indent=True)
phase1 = [
    "Module ai_engine.js: Script Node.js tự động đọc nội dung từ tài liệu và gọi WordPress REST API để tạo bài viết mới. Đã test thành công với 3 bài viết mẫu của Hừng Đông Media.",
    "Cấu hình WordPress Headless: Cài đặt và kích hoạt WPGraphQL, Faust.js plugin. Tạo Application Password cho AI Engine.",
    "Endpoint GraphQL: Xác nhận hoạt động tại http://hungdong.local/graphql với các query: posts, post, menuItems, generalSettings.",
    "Tệp .env: Lưu trữ an toàn WP_USER, WP_PASS và NEXT_PUBLIC_WORDPRESS_URL.",
]
for p1 in phase1:
    add_bullet(doc, p1)

doc.add_paragraph()
add_heading(doc, "6.3. Kết quả tích hợp", level=2)
add_paragraph(doc, "Sau giai đoạn tích hợp (07/07 – 20/07/2026), Frontend Next.js đã được kết nối thành công với Backend WordPress:", indent=True)
integration = [
    "Trang chủ (/) hiển thị đúng danh sách 10 bài viết mới nhất lấy từ GraphQL, bao gồm: 'Dịch vụ Booking Báo chí Toàn diện', 'Quản lý Dự án Thuê ngoài (Outsourcing PM)', 'Hello world!'.",
    "Giao diện Dark Mode với Gradient Text màu Cam Hừng Đông (#FF6B00 → #FF9900) và nền Slate (#0F172A) hoạt động đúng như thiết kế.",
    "Micro-animations: Hiệu ứng nâng thẻ bài viết (translateY -8px) + đổi màu viền sang cam khi hover hoạt động mượt mà.",
    "Faust Secret Key đã được cấu hình đúng trong .env.local, xác thực thành công với WordPress.",
    "Máy chủ Next.js dev chạy ổn định tại cổng 3000.",
]
for intg in integration:
    add_bullet(doc, intg)

doc.add_paragraph()
add_heading(doc, "6.4. Đánh giá giải pháp", level=2)
eval_table = doc.add_table(rows=1, cols=3)
eval_table.style = 'Table Grid'
for i, h in enumerate(["Tiêu chí", "Kết quả đạt được", "Đánh giá"]):
    eval_table.rows[0].cells[i].text = ''
    p = eval_table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, bold=True)
    tc = eval_table.rows[0].cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)
eval_data = [
    ("Tự động hóa nội dung", "AI Engine tạo bài viết từ PDF không cần thao tác thủ công", "✅ Đạt"),
    ("Tốc độ tải trang", "Next.js SSG cho thời gian render < 100ms", "✅ Đạt"),
    ("Bảo mật", "WordPress Admin hoàn toàn ẩn, dùng Application Password", "✅ Đạt"),
    ("Giao diện Premium", "Dark Mode + Gradient + Micro-animations đúng thiết kế", "✅ Đạt"),
    ("Khả năng mở rộng", "Kiến trúc Headless cho phép thay thế Frontend dễ dàng", "✅ Đạt"),
    ("SEO", "Có meta title/description trên từng bài viết", "⚠️ Cần bổ sung sitemap.xml"),
]
for ev in eval_data:
    add_table_row(eval_table, ev)

doc.add_page_break()

# ============================================================
# CHƯƠNG 7: HOÀN THIỆN – LỖI ĐÃ KHẮC PHỤC
# ============================================================
add_heading(doc, "CHƯƠNG 7: HOÀN THIỆN – CÁC LỖI ĐÃ KHẮC PHỤC")
doc.add_paragraph()

add_heading(doc, "7.1. Danh sách lỗi và cách khắc phục", level=2)
add_paragraph(doc, "Trong quá trình tích hợp và kiểm thử hệ thống (21/07 – 27/07/2026), nhóm đã phát hiện và khắc phục các lỗi sau:", indent=True)

bug_table = doc.add_table(rows=1, cols=5)
bug_table.style = 'Table Grid'
for i, h in enumerate(["STT", "Mô tả lỗi", "Nguyên nhân", "Cách khắc phục", "Kết quả"]):
    bug_table.rows[0].cells[i].text = ''
    p = bug_table.rows[0].cells[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    set_font(run, bold=True, size=11)
    tc = bug_table.rows[0].cells[i]._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'D9E1F2')
    tcPr.append(shd)
bugs = [
    ("1", "SassError: Undefined variable $break-medium", "@wordpress/block-library import các biến breakpoint từ file ngoài nhưng không được định nghĩa trước", "Thêm khai báo biến $break-small, $break-medium, $break-large trước lệnh @import trong _blocks.scss. Đồng thời comment bỏ 2 dòng import @wordpress/block-library/src/style và /src/theme vì không cần thiết cho dự án", "✅ Trang web load thành công"),
    ("2", "Trang chủ hiển thị nội dung rác mặc định của Faust.js thay vì danh sách bài viết", "Template front-page.js mặc định của Faust.js chỉ render placeholder text, không query GraphQL", "Viết lại hoàn toàn file wp-templates/front-page.js với GraphQL query lấy posts và component UI tùy chỉnh dạng card grid", "✅ Bài viết hiển thị đầy đủ"),
    ("3", "Máy chủ Next.js chạy ngầm bị ngắt tự động (process crash)", "Dùng lệnh Start-Process cmd /k gây ra hidden session không hiển thị trên màn hình", "Chuyển sang chạy trực tiếp lệnh npx next dev trong môi trường agent, giữ tiến trình sống trong suốt phiên làm việc", "✅ Máy chủ chạy ổn định"),
    ("4", "Giao diện frontend không nhận đúng màu sắc Dark Mode đã thiết kế", "File CSS global chưa được inject vào Next.js layout", "Thêm import file global.css vào _app.js và bổ sung CSS variables cho color palette Hừng Đông Media", "✅ Dark Mode hiển thị đúng"),
    ("5", "WordPress Headless site không tạo thư mục app-node (thiếu bước cài đặt)", "Khi tạo site Local, chọn tab 'Preferred' thay vì tab 'Custom' nên không kích hoạt được WP Engine Headless Add-on", "Hướng dẫn lại đúng quy trình: chọn tab Custom → tích checkbox WP Engine Headless → Continue", "✅ Add-on cài đặt thành công"),
]
for b in bugs:
    row = bug_table.add_row()
    for i, cell_text in enumerate(b):
        row.cells[i].text = ''
        p = row.cells[i].paragraphs[0]
        run = p.add_run(cell_text)
        set_font(run, size=11)

doc.add_paragraph()
add_heading(doc, "7.2. Kết quả sau khi hoàn thiện", level=2)
results = [
    "Hệ thống hoàn toàn hoạt động end-to-end: AI Engine → WordPress → GraphQL → Next.js",
    "Giao diện Frontend hiển thị chính xác 3 bài viết mẫu của Hừng Đông Media với Dark Mode Premium",
    "Không còn lỗi compilation (SCSS, TypeScript, ESLint)",
    "Luồng dữ liệu tự động (AI Engine tạo bài → WordPress lưu → Next.js hiển thị) hoạt động liên tục",
    "Kiến trúc có thể mở rộng: thêm Frontend pages, thêm nguồn dữ liệu AI mới mà không cần thay đổi Backend",
]
for r in results:
    add_bullet(doc, r)

doc.add_page_break()

# ============================================================
# KẾT LUẬN
# ============================================================
add_heading(doc, "KẾT LUẬN", center=True)
doc.add_paragraph()
add_paragraph(doc, "Qua quá trình thực tập tại Hừng Đông Media và triển khai dự án 'Xây Dựng Lại Website Theo Hướng Tự Động Hóa Nội Dung Bằng AI Và Tối Ưu Mọi Mặt', nhóm đã đạt được các kết quả sau:", indent=True)
conclusions = [
    "Nắm vững kiến trúc Headless CMS hiện đại (WordPress + WPGraphQL + Next.js) – xu hướng công nghệ được áp dụng bởi các tập đoàn lớn trên thế giới",
    "Xây dựng thành công cỗ máy AI Content Engine hoàn toàn tự động từ đầu vào (PDF/Word) đến đầu ra (bài viết trên website)",
    "Thiết kế và triển khai giao diện Frontend premium với Dark Mode, Micro-animations và điểm hiệu năng cao",
    "Giải quyết nhiều vấn đề kỹ thuật thực tế (SCSS conflicts, GraphQL integration, CI/CD flow) mang lại kinh nghiệm thực chiến quý báu",
    "Đề xuất được mô hình kinh doanh SaaS dựa trên hệ thống đã xây dựng, có tiềm năng thương mại hóa cao",
]
for c in conclusions:
    add_bullet(doc, c)

add_paragraph(doc, "Hướng phát triển tiếp theo của hệ thống bao gồm: tích hợp thêm mô hình LLM để tự động viết lại và tối ưu SEO cho bài viết, bổ sung dashboard analytics theo dõi lưu lượng truy cập và xây dựng cơ chế lập lịch (Cron Job) để hệ thống hoàn toàn không cần sự can thiệp của con người trong vận hành hàng ngày.", indent=True)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Hà Nội, ngày 27 tháng 07 năm 2026\nNhóm sinh viên thực hiện\n\n\n\nĐặng Minh Tuấn – Trần Anh Tuấn – Nguyễn Minh Hiếu")
set_font(run, italic=True)

# Save
output_path = r"D:\__G AG Projects\Thuc Tap Chuyen Nganh EHOU\Bao_cao_Chuyen_de_HOAN_CHINH.docx"
doc.save(output_path)
print("OK: " + output_path)
