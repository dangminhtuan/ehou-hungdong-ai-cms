# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page setup: A4, landscape-like via two columns or just single wide page ---
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2)

# ---- Helper functions ----
def para(doc, text='', align='center', bold=False, size=12, space_before=0, space_after=0, italic=False):
    p = doc.add_paragraph()
    p.alignment = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
        'justify':WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
    return p

def add_run(p, text, bold=False, size=12, underline=False, italic=False):
    run = p.add_run(text)
    run.bold      = bold
    run.italic    = italic
    run.underline = underline
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return run

def set_cell_text(cell, text, bold=False, size=11, align='center', italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = {
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
    }[align]
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

BLANK = '......................................................'
BLANK_SHORT = '.................'

# ============================================================
# HEADER
# ============================================================
para(doc, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', align='center', bold=True, size=13)
para(doc, 'Độc lập – Tự do – Hạnh phúc', align='center', bold=False, size=12)
para(doc, '----------o0o----------', align='center', size=11)
para(doc, '')

para(doc, 'ĐƠN XIN XÁC NHẬN CỦA ĐƠN VỊ THỰC TẬP', align='center', bold=True, size=14, space_before=4)
para(doc, '')

# Kính gửi
p = para(doc, align='left', size=12)
add_run(p, 'Kính gửi: ', bold=True, size=12)
add_run(p, 'Hừng Đông Media', bold=False, size=12, underline=True)

para(doc, '')

# Giới thiệu
p = para(doc, align='justify', size=12)
add_run(p, 'Chúng em gồm có ', size=12)
add_run(p, '03', bold=True, size=12)
add_run(p, ' thành viên:', size=12)

# --- Bảng thành viên ---
tbl = doc.add_table(rows=4, cols=4)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
col_widths = [Cm(1.2), Cm(6), Cm(3), Cm(3)]
for i, row in enumerate(tbl.rows):
    for j, cell in enumerate(row.cells):
        cell.width = col_widths[j]

# Header row
headers = ['STT', 'Họ và tên', 'Ngày sinh', 'Lớp']
for j, h in enumerate(headers):
    set_cell_text(tbl.rows[0].cells[j], h, bold=True, size=11, align='center')
    shade_cell(tbl.rows[0].cells[j], 'D9D9D9')

# 3 member rows - blank
for i in range(1, 4):
    set_cell_text(tbl.rows[i].cells[0], str(i), bold=False, size=11, align='center')
    set_cell_text(tbl.rows[i].cells[1], '', size=11, align='left')
    set_cell_text(tbl.rows[i].cells[2], '', size=11, align='center')
    set_cell_text(tbl.rows[i].cells[3], '', size=11, align='center')

para(doc, '')

# Khoa / Trường
p = para(doc, align='left', size=12)
add_run(p, 'Khoa: ', bold=True, size=12)
add_run(p, 'Công nghệ thông tin', size=12)

p = para(doc, align='left', size=12)
add_run(p, 'Trường: ', bold=True, size=12)
add_run(p, 'ĐH Mở Hà Nội', size=12)

para(doc, '')
para(doc, 'Thông tin thực tập:', align='left', bold=True, size=12)

def info_line(doc, label, value=''):
    p = para(doc, align='left', size=12)
    add_run(p, label, size=12)
    add_run(p, value if value else BLANK, size=12, underline=True)
    return p

info_line(doc, 'Thời gian thực tập từ ngày: ')
info_line(doc, 'Đến ngày: ')
info_line(doc, 'Vị trí thực tập: ')
info_line(doc, 'Người hướng dẫn: ')
info_line(doc, 'Chức vụ: ')
info_line(doc, 'Email: ')
info_line(doc, 'Số điện thoại: ')
info_line(doc, 'Tại doanh nghiệp: ', 'Hừng Đông Media')

para(doc, '')

# Body text
p = para(doc, align='justify', size=12)
add_run(p, 'Nay chúng em làm đơn này kính trình Ban lãnh đạo Quý Cơ quan/Doanh nghiệp xác nhận cho nhóm chúng em về việc thực tập tại đơn vị trong khoảng thời gian nêu trên.', size=12)

para(doc, '')

p = para(doc, align='justify', size=12)
add_run(p, 'Trong thời gian thực tập, chúng em kính mong Quý Cơ quan/Doanh nghiệp tạo điều kiện để cả nhóm hoàn thành tốt nội dung thực tập và các nhiệm vụ được giao.', size=12)

p = para(doc, align='justify', size=12)
add_run(p, 'Chúng em xin chân thành cảm ơn!', size=12)

para(doc, '')

# Date line
p = para(doc, align='right', size=12, italic=True)
add_run(p, 'TP. Hà Nội, ngày ...... tháng ...... năm 2026', size=12, italic=True)

para(doc, '')

# --- Signature table: Đại diện ĐV thực tập | Đại diện nhóm SV ---
sig1 = doc.add_table(rows=5, cols=2)
sig1.alignment = WD_TABLE_ALIGNMENT.CENTER

# Remove borders
from docx.oxml import OxmlElement
def remove_borders(tbl):
    tblPr = tbl._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)

remove_borders(sig1)

set_cell_text(sig1.rows[0].cells[0], 'ĐẠI DIỆN ĐƠN VỊ THỰC TẬP', bold=True, size=11, align='center')
set_cell_text(sig1.rows[0].cells[1], 'ĐẠI DIỆN NHÓM SINH VIÊN', bold=True, size=11, align='center')
set_cell_text(sig1.rows[1].cells[0], '(Ký, ghi rõ họ tên, đóng dấu xác nhận)', bold=False, size=10, align='center')
set_cell_text(sig1.rows[1].cells[1], '(Ký và ghi rõ họ tên)', bold=False, size=10, align='center')
# Blank rows for signature space
for i in [2, 3]:
    set_cell_text(sig1.rows[i].cells[0], '', size=11, align='center')
    set_cell_text(sig1.rows[i].cells[1], '', size=11, align='center')
set_cell_text(sig1.rows[4].cells[0], '', size=11, align='center')
set_cell_text(sig1.rows[4].cells[1], 'Đặng Minh Tuấn', size=11, align='center')

para(doc, '')
para(doc, 'CHỮ KÝ CỦA CÁC THÀNH VIÊN TRONG NHÓM', align='center', bold=True, size=12)
para(doc, '')

# --- 3-member signature table ---
sig2 = doc.add_table(rows=4, cols=3)
sig2.alignment = WD_TABLE_ALIGNMENT.CENTER
remove_borders(sig2)

members = ['THÀNH VIÊN 1', 'THÀNH VIÊN 2', 'THÀNH VIÊN 3']
member_names = ['Đặng Minh Tuấn', 'Trần Anh Tuấn', 'Nguyễn Minh Hiếu']

for j, (m, n) in enumerate(zip(members, member_names)):
    set_cell_text(sig2.rows[0].cells[j], m, bold=True, size=11, align='center')
    set_cell_text(sig2.rows[1].cells[j], '(Ký và ghi rõ họ tên)', bold=False, size=10, align='center')

# Blank space for signature (rows 2)
for j in range(3):
    set_cell_text(sig2.rows[2].cells[j], '', size=11)
    set_cell_text(sig2.rows[2].cells[j], '', size=11)

# Name row
for j, n in enumerate(member_names):
    set_cell_text(sig2.rows[3].cells[j], n, bold=False, size=11, align='center')

# Save
out = r"D:\__G AG Projects\Thuc Tap Chuyen Nganh EHOU\Don_Xin_Xac_Nhan_Thuc_Tap.docx"
doc.save(out)
print("OK: " + out)
