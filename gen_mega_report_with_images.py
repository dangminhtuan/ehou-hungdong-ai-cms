# -*- coding: utf-8 -*-
import sys
import os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

FIGURES_DIR = r"D:\__G AG Projects\Thuc Tap Chuyen Nganh EHOU\report_figures"

# --- Page margins (Chuẩn format đồ án EHOU: Trái 3cm, Phải 2cm, Trên 2.5cm, Dưới 2.5cm) ---
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

# --- Base Styles ---
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(13)

def set_font(run, bold=False, size=13, color=None, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, bold=True, size=14)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, bold=True, size=13)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, bold=True, italic=True, size=13)
    return p

def add_paragraph(doc, text, indent=True, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, size=13)
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1.27)
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, bold=True, size=13)
    r_txt = p.add_run(text)
    set_font(r_txt, size=13)
    return p

def add_figure(doc, img_name, caption_text, width_inch=5.8):
    img_path = os.path.join(FIGURES_DIR, img_name)
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(10)
        p_img.paragraph_format.space_after = Pt(4)
        run_img = p_img.add_run()
        run_img.add_picture(img_path, width=Inches(width_inch))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(2)
        p_cap.paragraph_format.space_after = Pt(12)
        run_cap = p_cap.add_run(caption_text)
        set_font(run_cap, bold=True, italic=True, size=11, color=(75, 85, 99))
    else:
        print(f"Warning: image {img_path} not found.")

def format_table_header(row, col_widths, headers):
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.width = col_widths[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(h)
        set_font(run, bold=True, size=11)
        # Shading
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9E1F2')
        tcPr.append(shd)

def add_table_row(table, col_widths, cells_data, is_center_list=None):
    row = table.add_row()
    for i, data in enumerate(cells_data):
        cell = row.cells[i]
        cell.width = col_widths[i]
        cell.text = ''
        p = cell.paragraphs[0]
        if is_center_list and is_center_list[i]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(str(data))
        set_font(run, size=11)
    return row

# ============================================================
# TRANG BÌA CHÍNH THỨC
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TRƯỜNG ĐẠI HỌC MỞ HÀ NỘI\nVIỆN ĐÀO TẠO & PHÁT TRIỂN HỌC TẬP SUỐT ĐỜI\nKHOA CÔNG NGHỆ THÔNG TIN")
set_font(run, bold=True, size=14)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BÁO CÁO CHUYÊN ĐỀ THỰC TẬP CHUYÊN NGÀNH\n(IT43.027 - IT43.028)")
set_font(run, bold=True, size=16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ĐỀ TÀI:")
set_font(run, bold=True, size=14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("XÂY DỰNG LẠI WEBSITE THEO HƯỚNG TỰ ĐỘNG HÓA NỘI DUNG\nBẰNG AI VÀ TỐI ƯU MỌI MẶT\n(MÔ HÌNH HEADLESS WORDPRESS KẾT HỢP NEXT.JS VÀ AI CONTENT ENGINE)")
set_font(run, bold=True, size=15)

doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=5, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_data = [
    ("Đơn vị thực tập:", "Hừng Đông Media Solutions"),
    ("Người hướng dẫn doanh nghiệp:", "Ban Giám Đốc Hừng Đông Media"),
    ("Nhóm sinh viên thực hiện:", "1. Đặng Minh Tuấn   - Lớp: CHTM518 (Nhóm trưởng)\n2. Trần Anh Tuấn    - Lớp: CHCT419\n3. Nguyễn Minh Hiếu - Lớp: CLCA520"),
    ("Chuyên ngành đào tạo:", "Công nghệ Thông tin"),
    ("Niên khóa thực tập:", "2025 – 2026 (Tuần 1 đến Tuần 16)"),
]
for i, (label, val) in enumerate(info_data):
    cell_lbl, cell_val = info_table.rows[i].cells
    cell_lbl.text = ''
    cell_val.text = ''
    p1 = cell_lbl.paragraphs[0]
    p2 = cell_val.paragraphs[0]
    r1 = p1.add_run(label)
    r2 = p2.add_run(val)
    set_font(r1, bold=True, size=12)
    set_font(r2, size=12)

# Remove borders info_table
for row in info_table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ['top','left','bottom','right']:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            tcBorders.append(b)
        tcPr.append(tcBorders)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("HÀ NỘI – NĂM 2026")
set_font(run, bold=True, size=13)

doc.add_page_break()

# ============================================================
# LỜI CẢM ƠN VÀ CAM KẾT
# ============================================================
add_heading_1(doc, "LỜI CẢM ƠN")
add_paragraph(doc, "Lời đầu tiên, nhóm sinh viên xin gửi lời cảm ơn chân thành và sâu sắc nhất đến Ban Giám hiệu Trường Đại học Mở Hà Nội, Viện Đào tạo & Phát triển Học tập Suốt đời, cùng toàn thể quý thầy cô giáo Khoa Công nghệ Thông tin đã tận tình truyền đạt những kiến thức chuyên môn quý báu và định hướng phương pháp nghiên cứu học thuật trong suốt quá trình học tập.", indent=True)
add_paragraph(doc, "Đặc biệt, nhóm xin gửi lời cảm ơn trân trọng đến Giảng viên hướng dẫn môn học TS. Vũ Xuân Hạnh và Quản lý lớp học phần ThS. Nguyễn Hữu Toàn đã luôn đồng hành, góp ý và tạo điều kiện thuận lợi nhất để nhóm hoàn thành các mốc báo cáo tiến độ theo đúng kế hoạch chuyên đề thực tập chuyên ngành.", indent=True)
add_paragraph(doc, "Đồng thời, nhóm xin chân thành cảm ơn Ban Lãnh đạo và các anh/chị đồng nghiệp tại Công ty Cổ phần Truyền thông Hừng Đông (Hừng Đông Media) đã tiếp nhận, tạo môi trường làm việc thực tế, cung cấp tài liệu nghiệp vụ và hỗ trợ nhóm thử nghiệm cỗ máy tự động hóa nội dung bằng AI trên hạ tầng website của doanh nghiệp.", indent=True)

add_heading_1(doc, "LỜI CAM ĐOAN")
add_paragraph(doc, "Chúng tôi xin cam đoan bản báo cáo chuyên đề thực tập 'Xây dựng lại website theo hướng tự động hóa nội dung bằng AI và tối ưu mọi mặt' là công trình nghiên cứu và phát triển sản phẩm thực tế của nhóm chúng tôi dưới sự hướng dẫn của giảng viên và doanh nghiệp tiếp nhận thực tập.", indent=True)
add_paragraph(doc, "Toàn bộ các số liệu khảo sát, kiến trúc giải pháp, mã nguồn chương trình (Node.js AI Engine, Next.js Frontend, WordPress Headless CMS) và các kết quả kiểm thử thực tế được trình bày trong báo cáo là trung thực, minh bạch và không sao chép trái phép từ bất kỳ công trình nào khác.", indent=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Hà Nội, ngày 01 tháng 09 năm 2026\nĐại diện nhóm sinh viên thực hiện\n\n\n\nĐặng Minh Tuấn")
set_font(run, bold=True, italic=True)

doc.add_page_break()

# ============================================================
# MỤC LỤC TỔNG THỂ & DANH MỤC HÌNH ẢNH
# ============================================================
add_heading_1(doc, "DANH MỤC HÌNH ẢNH MINH HỌA")
figs_list = [
    ("Hình 1.1: Sơ đồ Cơ cấu tổ chức & Quy trình Agile tại Hừng Đông Media", "5"),
    ("Hình 2.1: Hiệu năng website cũ đo bằng Google Lighthouse (Báo động đỏ 32/100)", "7"),
    ("Hình 2.2: Sơ đồ Use Case tổng thể của hệ thống", "9"),
    ("Hình 3.1: Kiến trúc phân tách 3 tầng Decoupled Headless CMS", "11"),
    ("Hình 3.2: Sơ đồ thực thể cơ sở dữ liệu (Database ERD Schema)", "13"),
    ("Hình 4.1: Khung mã nguồn xử lý cỗ máy AI Content Engine (ai_engine.js)", "15"),
    ("Hình 4.2: Khung mã nguồn truy vấn GraphQL trang chủ (front-page.js)", "16"),
    ("Hình 5.1: Khung mã nguồn giải quyết xung đột SCSS Breakpoints (_blocks.scss)", "18"),
    ("Hình 5.2: Kết quả kiểm thử Google Lighthouse đạt điểm tuyệt đối 100/100", "20"),
    ("Hình 6.1: Infographic mô phỏng Định luật Tăng trưởng GAS (G = A² × S)", "22"),
    ("Hình 6.2: Mô hình Trạm nạp nhiên liệu số GAS Fueling Station cho doanh nghiệp", "23"),
    ("Hình 6.3: Cấu trúc tệp chuẩn hóa llms.txt phục vụ Generative Engine Optimization", "24"),
    ("Hình 7.1: Sơ đồ Gantt Chart tiến độ thực hiện đồ án 16 tuần", "26"),
]
for item, pg in figs_list:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(item)
    set_font(run, size=12)

doc.add_page_break()

# ============================================================
# NỘI DUNG CHƯƠNG 1
# ============================================================
add_heading_1(doc, "CHƯƠNG 1: TỔNG QUAN DOANH NGHIỆP VÀ VỊ TRÍ THỰC TẬP (TUẦN 1 – 3)")

add_heading_2(doc, "1.1. Giới thiệu tổng quan về Hừng Đông Media")
add_paragraph(doc, "Công ty Cổ phần Truyền thông Hừng Đông (Hừng Đông Media) là đơn vị hoạt động chuyên sâu trong lĩnh vực tư vấn chiến lược truyền thông, cung cấp giải pháp tiếp thị kỹ thuật số (Digital Marketing), quan hệ công chúng (Public Relations – PR) và phát triển các nền tảng công nghệ phục vụ doanh nghiệp vừa và nhỏ (SMBs) tại Việt Nam.")
add_paragraph(doc, "Trải qua quá trình hình thành và phát triển, Hừng Đông Media đã xây dựng mạng lưới đối tác sâu rộng với hơn 200 cơ quan báo chí, trang tin điện tử hàng đầu cả nước (như VnExpress, Dân Trí, CafeF, VietnamNet, Tuổi Trẻ...). Doanh nghiệp định vị mình là cầu nối chiến lược giúp khách hàng nâng tầm nhận diện thương hiệu, tối ưu hóa chi phí quảng cáo và tiếp cận đúng đối tượng khách hàng mục tiêu.")

add_heading_2(doc, "1.2. Cơ cấu tổ chức và quy trình nghiệp vụ")
add_paragraph(doc, "Cơ cấu tổ chức của Hừng Đông Media được tinh gọn hóa theo mô hình Agile, bao gồm 4 khối phòng ban chính:")
add_bullet(doc, "Khối Điều Hành & Chiến Lược (Board of Directors): Chịu trách nhiệm hoạch định chiến lược kinh doanh, định hướng phát triển sản phẩm công nghệ và kiểm soát tài chính.", "1. ")
add_bullet(doc, "Khối Công Nghệ & Giải Pháp Số (Tech & Digital Solutions): Trực tiếp nghiên cứu, lập trình và vận hành các nền tảng website, ứng dụng AI và hạ tầng máy chủ đám mây.", "2. ")
add_bullet(doc, "Khối Truyền Thông & Sáng Tạo Nội Dung (Media & Creative): Quản lý mạng lưới quan hệ báo chí, sản xuất tài liệu hồ sơ năng lực, kịch bản PR và chiến dịch truyền thông.", "3. ")
add_bullet(doc, "Khối Kinh Doanh & Dịch Vụ Khách Hàng (Account & Sales): Chăm sóc đối tác, tư vấn giải pháp, đàm phán hợp đồng và hỗ trợ khách hàng sau bán hàng.", "4. ")

add_figure(doc, "fig_1_1_org_chart.png", "Hình 1.1: Sơ đồ Cơ cấu tổ chức & Vị trí thực tập tại Hừng Đông Media")

add_heading_2(doc, "1.3. Vị trí thực tập và phân công nhiệm vụ thành viên")
add_paragraph(doc, "Nhóm sinh viên thực tập thuộc Trường Đại học Mở Hà Nội được tiếp nhận vào Khối Công Nghệ & Giải Pháp Số với vị trí Kỹ sư Phát triển Hệ thống (Software Engineer Intern). Dưới sự hướng dẫn của Mentor doanh nghiệp, nhóm chịu trách nhiệm toàn diện trong việc khảo sát, thiết kế kiến trúc, lập trình cỗ máy AI Engine và xây dựng Frontend Next.js.")

table_assign = doc.add_table(rows=1, cols=4)
table_assign.style = 'Table Grid'
table_assign.alignment = WD_TABLE_ALIGNMENT.CENTER
widths_assign = [Cm(1.2), Cm(4.0), Cm(6.5), Cm(3.5)]
format_table_header(table_assign.rows[0], widths_assign, ["STT", "Họ và tên", "Nhiệm vụ chuyên môn đảm nhiệm", "Vai trò nhóm"])

assign_data = [
    ("1", "Đặng Minh Tuấn\n(CHTM518)", "Khảo sát yêu cầu, thiết kế kiến trúc tổng thể Headless CMS, lập trình cỗ máy AI Content Engine (ai_engine.js, auto_cron.js), tích hợp REST API và chịu trách nhiệm chung về tiến độ dự án.", "Nhóm trưởng\n(Lead Dev)"),
    ("2", "Trần Anh Tuấn\n(CHCT419)", "Phân tích đặc tả Use Case, thiết kế cơ sở dữ liệu WordPress, cấu hình schema GraphQL, tham gia lập trình module xử lý trích xuất văn bản và viết tài liệu kiểm thử.", "Thành viên\n(Backend Dev)"),
    ("3", "Nguyễn Minh Hiếu\n(CLCA520)", "Thiết kế giao diện người dùng (UI/UX), lập trình Frontend Next.js 14 với Faust.js, cấu hình SCSS Dark Mode, tối ưu hóa điểm Google Lighthouse và viết kịch bản hướng dẫn vận hành.", "Thành viên\n(Frontend Dev)"),
]
for row_d in assign_data:
    add_table_row(table_assign, widths_assign, row_d, [True, False, False, True])

doc.add_page_break()

# ============================================================
# NỘI DUNG CHƯƠNG 2
# ============================================================
add_heading_1(doc, "CHƯƠNG 2: KHẢO SÁT HIỆN TRẠNG VÀ PHÁT BIỂU BÀI TOÁN (TUẦN 4 – 5)")

add_heading_2(doc, "2.1. Khảo sát hiện trạng và các nút thắt của website truyền thống")
add_paragraph(doc, "Khảo sát thực tế tại Hừng Đông Media và hơn 30 website doanh nghiệp đối tác cho thấy hầu hết các website hiện nay đều được xây dựng dựa trên kiến trúc Monolithic WordPress truyền thống. Điểm số đo lường hiệu năng bằng công cụ Google Lighthouse trên các website này thường ở mức báo động đỏ (chỉ đạt từ 25 – 45 điểm):")

add_figure(doc, "fig_2_1_lighthouse_old.png", "Hình 2.1: Hiệu năng website cũ đo bằng Google Lighthouse (Báo động đỏ 32/100)")

add_paragraph(doc, "Các nút thắt (bottlenecks) nghiêm trọng cản trở sự phát triển của doanh nghiệp bao gồm:")
add_bullet(doc, "Tốc độ tải trang chậm và điểm Core Web Vitals thấp: Thời gian tải trang kéo dài từ 4 đến 7 giây, khiến hơn 53% khách hàng tiềm năng thoát trang ngay lập tức.", "• ")
add_bullet(doc, "Rủi ro bảo mật và nguy cơ bị tấn công: Cổng đăng nhập WordPress Admin và cơ sở dữ liệu nằm trực tiếp trên máy chủ công khai, thường xuyên trở thành mục tiêu của các cuộc tấn công mã độc.", "• ")
add_bullet(doc, "Chi phí nhân sự sản xuất nội dung quá lớn: Doanh nghiệp muốn duy trì thứ hạng SEO phải thuê đội ngũ nhân viên Content Marketing viết bài thủ công với chi phí từ 10 – 15 triệu đồng/tháng.", "• ")

add_heading_2(doc, "2.2. Sơ đồ Use Case và phân tích tác nhân")
add_paragraph(doc, "Hệ thống được thiết kế phục vụ 3 nhóm tác nhân chính (Doanh nghiệp, Cỗ máy AI và Người dùng cuối/Google Bot):")

add_figure(doc, "fig_2_2_use_case.png", "Hình 2.2: Sơ đồ Use Case tổng thể của hệ thống")

doc.add_page_break()

# ============================================================
# NỘI DUNG CHƯƠNG 3
# ============================================================
add_heading_1(doc, "CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG (TUẦN 6 – 7)")

add_heading_2(doc, "3.1. Kiến trúc phân tách Headless CMS (Decoupled Architecture)")
add_paragraph(doc, "Hệ thống áp dụng mô hình kiến trúc Decoupled 3 tầng hiện đại, tách biệt hoàn toàn giữa Tầng Dữ liệu (Backend Core), Tầng Xử lý AI (Automation Engine) và Tầng Hiển thị (Frontend Next.js):")

add_figure(doc, "fig_3_1_architecture.png", "Hình 3.1: Kiến trúc phân tách 3 tầng (Decoupled Headless CMS)")

add_heading_2(doc, "3.2. Thiết kế Cơ sở dữ liệu và Data Schema")
add_paragraph(doc, "Cơ sở dữ liệu của hệ thống kế thừa cấu trúc chuẩn hóa cao của WordPress MySQL kết hợp với các Custom Post Types và GraphQL Schema mở rộng:")

add_figure(doc, "fig_3_2_database_erd.png", "Hình 3.2: Sơ đồ thực thể cơ sở dữ liệu (Database ERD Schema)")

doc.add_page_break()

# ============================================================
# NỘI DUNG CHƯƠNG 4
# ============================================================
add_heading_1(doc, "CHƯƠNG 4: LẬP TRÌNH VÀ TÍCH HỢP HỆ THỐNG (TUẦN 8 – 10)")

add_heading_2(doc, "4.1. Lập trình Cỗ máy AI Content Engine (ai_engine.js)")
add_paragraph(doc, "Cỗ máy AI Content Engine được lập trình bằng Node.js đặt trong thư mục `codehungdong/`. Đoạn mã nguồn xử lý việc đọc dữ liệu và gửi HTTP POST qua REST API được thiết kế như sau:")

add_figure(doc, "fig_4_5_code_ai_engine.png", "Hình 4.1: Khung mã nguồn xử lý cỗ máy AI Content Engine (ai_engine.js)")

add_heading_2(doc, "4.2. Xây dựng Giao diện Frontend Next.js 14 Premium Dark Mode")
add_paragraph(doc, "Trang chủ (`front-page.js`) sử dụng Apollo Client để truy vấn GraphQL lấy 10 bài viết mới nhất và hiển thị dạng thẻ Grid sang trọng:")

add_figure(doc, "fig_4_6_code_front_page.png", "Hình 4.2: Khung mã nguồn truy vấn GraphQL trang chủ (front-page.js)")

doc.add_page_break()

# ============================================================
# NỘI DUNG CHƯƠNG 5
# ============================================================
add_heading_1(doc, "CHƯƠNG 5: KIỂM THỬ VÀ KHẮC PHỤC SỰ CỐ KỸ THUẬT (TUẦN 11 – 12)")

add_heading_2(doc, "5.1. Khắc phục sự cố biên dịch SCSS Breakpoints")
add_paragraph(doc, "Trong quá trình tích hợp, lỗi `SassError: Undefined variable $break-medium` đã được khắc phục triệt để bằng cách định nghĩa bổ sung các biến breakpoint toàn cục trong file `_blocks.scss`:")

add_figure(doc, "fig_5_2_code_bugfix_scss.png", "Hình 5.1: Khung mã nguồn giải quyết xung đột SCSS Breakpoints (_blocks.scss)")

add_heading_2(doc, "5.2. Kết quả kiểm định hiệu năng Google Lighthouse 100/100")
add_paragraph(doc, "Sau khi tối ưu hóa toàn diện trên nền tảng Next.js 14 Static Site Generation (SSG), hệ thống đạt điểm số tuyệt đối 100/100 trên cả 4 hạng mục đánh giá của Google:")

add_figure(doc, "fig_5_3_lighthouse_100.png", "Hình 5.2: Kết quả kiểm thử Google Lighthouse đạt điểm tuyệt đối 100/100")

doc.add_page_break()

# ============================================================
# NỘI DUNG CHƯƠNG 6
# ============================================================
add_heading_1(doc, "CHƯƠNG 6: TRIỂN KHAI VÀ PHƯƠNG ÁN THƯƠNG MẠI HÓA (TUẦN 13 – 14)")

add_heading_2(doc, "6.1. Định luật Tăng trưởng GAS (Growth = AI² × Speed)")
add_paragraph(doc, "Trên hệ thống thương mại, chúng tôi đã phát triển Định luật Tăng trưởng Độc quyền GAS mô phỏng phương trình $E = mc^2$:")

add_figure(doc, "fig_6_2_gas_formula.png", "Hình 6.1: Infographic mô phỏng Định luật Tăng trưởng GAS (G = A² × S)")

add_heading_2(doc, "6.2. Mô hình Trạm Nạp Nhiên Liệu Số (GAS Fueling Station)")
add_paragraph(doc, "Hình tượng Trạm nạp nhiên liệu số giúp doanh nghiệp hình dung trực quan về dịch vụ tự động hóa nội dung và gia tốc website:")

add_figure(doc, "fig_6_3_gas_station.jpg", "Hình 6.2: Mô hình Trạm nạp nhiên liệu số GAS Fueling Station cho doanh nghiệp")

add_heading_2(doc, "6.3. Tiêu chuẩn Generative Engine Optimization (GEO) & llms.txt")
add_paragraph(doc, "Hệ thống hỗ trợ tệp chuẩn hóa `public/llms.txt` giúp các mô hình AI thế hệ mới (ChatGPT, Perplexity, Gemini) trích xuất dữ liệu công ty trong 0.05 giây:")

add_figure(doc, "fig_6_5_llms_txt.png", "Hình 6.3: Cấu trúc tệp chuẩn hóa llms.txt phục vụ Generative Engine Optimization")

doc.add_page_break()

# ============================================================
# NỘI DUNG CHƯƠNG 7
# ============================================================
add_heading_1(doc, "CHƯƠNG 7: TỔNG KẾT VÀ ĐÓNG GÓI HỒ SƠ THỰC TẬP (TUẦN 15 – 16)")

add_heading_2(doc, "7.1. Sơ đồ Gantt Chart tiến độ 16 tuần")
add_paragraph(doc, "Toàn bộ quá trình thực tập trải dài qua 16 tuần được tổng hợp trực quan qua sơ đồ tiến độ sau:")

add_figure(doc, "fig_7_1_gantt_chart.png", "Hình 7.1: Sơ đồ Gantt Chart tiến độ thực hiện đồ án 16 tuần")

add_heading_2(doc, "7.2. Kết luận và định hướng phát triển")
add_paragraph(doc, "Đề tài 'Xây dựng lại website theo hướng tự động hóa nội dung bằng AI và tối ưu mọi mặt' đã hoàn thành xuất sắc 100% mục tiêu chuyên môn và chứng minh được tiềm năng kinh doanh thực tế, sẵn sàng đưa vào vận hành thương mại tại Hừng Đông Media.")

# Save final mega report
output_path = r"D:\__G AG Projects\Thuc Tap Chuyen Nganh EHOU\Bao_cao_Chuyen_de_HOAN_CHINH_16_Tuan_FULL_HINH_ANH.docx"
doc.save(output_path)
print("SUCCESS: " + output_path)
